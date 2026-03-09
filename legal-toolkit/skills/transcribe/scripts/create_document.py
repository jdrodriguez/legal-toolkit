#!/usr/bin/env python3
"""
Generate a professional .docx transcript document.

Reads transcript data from a work directory and produces a formatted Word
document with metadata, executive summary, topics, full transcript, and quotes.

Usage:
    python3 create_document.py <work_dir> <output_path> [--analysis <analysis_json>]

Arguments:
    work_dir     Path to the transcript work directory (contains transcript.txt, metadata.json)
    output_path  Path for the output .docx file

Options:
    --analysis   JSON string with synthesized analysis (executive_summary, key_topics,
                 action_items, notable_quotes). If not provided, generates minimal analysis.
"""
import argparse
import json
import os
import re
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def read_file(path: str) -> str:
    """Read a text file, return empty string if missing."""
    try:
        with open(path, encoding="utf-8") as f:
            return f.read()
    except (FileNotFoundError, OSError):
        return ""


def read_json(path: str) -> dict:
    """Read a JSON file, return empty dict if missing."""
    try:
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError, OSError):
        return {}


def format_duration(seconds: float) -> str:
    """Convert seconds to HH:MM:SS."""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    if h > 0:
        return f"{h}:{m:02d}:{s:02d}"
    return f"{m}:{s:02d}"


def parse_transcript_lines(text: str) -> list[dict]:
    """Parse transcript text into structured entries.

    Handles both formats from the worker:
      [HH:MM:SS - HH:MM:SS] SPEAKER_00: text   (with end timestamp and diarization)
      [HH:MM:SS - HH:MM:SS] text                (with end timestamp, no speaker)
      [HH:MM:SS] SPEAKER: text                   (start timestamp only)
      [HH:MM:SS] text                            (start timestamp only, no speaker)
    """
    entries = []
    # Step 1: extract timestamp bracket and remainder
    ts_pattern = re.compile(r"^\[(\d{1,2}:\d{2}:\d{2})[^\]]*\]\s*(.+)$")
    # Step 2: check if remainder starts with a speaker label
    speaker_pattern = re.compile(r"^([A-Za-z][A-Za-z0-9_ ]*?):\s*(.+)$")

    for line in text.strip().splitlines():
        line = line.strip()
        if not line:
            continue
        ts_match = ts_pattern.match(line)
        if ts_match:
            timestamp = ts_match.group(1)
            remainder = ts_match.group(2)
            sp_match = speaker_pattern.match(remainder)
            if sp_match:
                speaker = sp_match.group(1)
                body = sp_match.group(2)
            else:
                speaker = ""
                body = remainder
            entries.append({
                "timestamp": timestamp,
                "speaker": speaker,
                "text": body,
            })
        else:
            # Non-timestamped line — append to previous entry or add as-is
            if entries:
                entries[-1]["text"] += " " + line
            else:
                entries.append({"timestamp": "", "speaker": "", "text": line})
    return entries


# ---------------------------------------------------------------------------
# Document styling
# ---------------------------------------------------------------------------
def setup_styles(doc: Document):
    """Configure document styles for professional appearance."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading styles
    for level, size, color in [
        ("Heading 1", 22, RGBColor(0x1A, 0x1A, 0x2E)),
        ("Heading 2", 16, RGBColor(0x2D, 0x2D, 0x44)),
        ("Heading 3", 13, RGBColor(0x44, 0x44, 0x66)),
    ]:
        h = doc.styles[level]
        h.font.name = "Arial"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = color
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)

    # Create transcript style
    try:
        ts = doc.styles.add_style("Transcript", WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        ts = doc.styles["Transcript"]
    ts.font.name = "Consolas"
    ts.font.size = Pt(9.5)
    ts.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    ts.paragraph_format.space_after = Pt(3)
    ts.paragraph_format.line_spacing = 1.3

    return doc


def add_metadata_table(doc: Document, metadata: dict, transcript_text: str):
    """Add a metadata summary table."""
    duration = metadata.get("duration_seconds", 0)
    language = metadata.get("language_detected", metadata.get("language", "auto"))
    model = metadata.get("model_used", metadata.get("model", "unknown"))
    speakers = metadata.get("speaker_count", metadata.get("num_speakers", 0))
    segments = metadata.get("segment_count", metadata.get("num_segments", 0))
    word_count = metadata.get("word_count", len(transcript_text.split()))

    rows = [
        ("Duration", format_duration(duration)),
        ("Language", language.upper() if language != "auto" else "Auto-detected"),
        ("Model", model),
        ("Speakers", str(speakers) if speakers > 0 else "Not identified"),
        ("Segments", str(segments)),
        ("Word Count", f"{word_count:,}"),
        ("Transcribed", datetime.now().strftime("%B %d, %Y")),
    ]

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(rows):
        cell_label = table.rows[i].cells[0]
        cell_value = table.rows[i].cells[1]
        cell_label.text = label
        cell_value.text = value
        # Bold the labels
        for paragraph in cell_label.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
        for paragraph in cell_value.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

    doc.add_paragraph("")  # spacer


def add_speaker_stats(doc: Document, metadata: dict):
    """Add speaker statistics table if diarization data is available."""
    stats = metadata.get("speakers", metadata.get("speaker_stats", {}))
    if not stats:
        return

    doc.add_heading("Speaker Statistics", level=1)

    headers = ["Speaker", "Segments", "Words", "Duration", "% of Total"]
    table = doc.add_table(rows=1 + len(stats), cols=5)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    total_duration = sum(s.get("total_seconds", s.get("duration_seconds", 0)) for s in stats.values())
    for i, (speaker, data) in enumerate(sorted(stats.items()), start=1):
        dur = data.get("total_seconds", data.get("duration_seconds", 0))
        pct = (dur / total_duration * 100) if total_duration > 0 else 0
        values = [
            speaker,
            str(data.get("segment_count", 0)),
            str(data.get("word_count", 0)),
            format_duration(dur),
            f"{pct:.1f}%",
        ]
        for j, val in enumerate(values):
            cell = table.rows[i].cells[j]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph("")


def add_transcript_body(doc: Document, entries: list[dict]):
    """Add the full transcript with timestamps and speaker labels."""
    doc.add_heading("Full Transcript", level=1)

    current_speaker = None
    for entry in entries:
        ts = entry["timestamp"]
        speaker = entry["speaker"]
        text = entry["text"]

        p = doc.add_paragraph(style="Transcript")

        # Timestamp
        if ts:
            run_ts = p.add_run(f"[{ts}]  ")
            run_ts.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run_ts.font.size = Pt(9)

        # Speaker label (only show when it changes)
        if speaker and speaker != current_speaker:
            run_sp = p.add_run(f"{speaker}: ")
            run_sp.bold = True
            run_sp.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            run_sp.font.size = Pt(10)
            current_speaker = speaker

        # Text
        run_text = p.add_run(text)
        run_text.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Main document builder
# ---------------------------------------------------------------------------
def build_document(
    work_dir: str,
    output_path: str,
    analysis: dict | None = None,
) -> str:
    """Build the transcript .docx document.

    Args:
        work_dir: Path to the work directory with transcript.txt and metadata.json.
        output_path: Path to write the output .docx file.
        analysis: Optional dict with keys: executive_summary, key_topics,
                  action_items, notable_quotes. If None, minimal analysis is used.

    Returns:
        The absolute path to the created .docx file.
    """
    metadata = read_json(os.path.join(work_dir, "metadata.json"))
    transcript_text = read_file(os.path.join(work_dir, "transcript.txt"))

    if not transcript_text.strip():
        raise ValueError(f"No transcript found at {work_dir}/transcript.txt")

    source_file = metadata.get("source_file", "Unknown")
    filename = os.path.basename(source_file)

    doc = Document()
    setup_styles(doc)

    # --- Title ---
    title = doc.add_heading(f"Transcript: {filename}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Style the title
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        run.font.size = Pt(26)

    doc.add_paragraph("")  # spacer

    # --- Metadata Table ---
    add_metadata_table(doc, metadata, transcript_text)

    # --- Executive Summary ---
    if analysis and analysis.get("executive_summary"):
        doc.add_heading("Executive Summary", level=1)
        for para_text in analysis["executive_summary"].split("\n\n"):
            para_text = para_text.strip()
            if para_text:
                doc.add_paragraph(para_text)

    # --- Key Topics ---
    if analysis and analysis.get("key_topics"):
        doc.add_heading("Key Topics", level=1)
        for topic in analysis["key_topics"]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(topic)

    # --- Action Items ---
    if analysis and analysis.get("action_items"):
        doc.add_heading("Action Items", level=1)
        for item in analysis["action_items"]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item)

    # --- Speaker Statistics ---
    add_speaker_stats(doc, metadata)

    # --- Full Transcript ---
    entries = parse_transcript_lines(transcript_text)
    add_transcript_body(doc, entries)

    # --- Notable Quotes ---
    if analysis and analysis.get("notable_quotes"):
        doc.add_heading("Notable Quotes", level=1)
        for quote in analysis["notable_quotes"]:
            p = doc.add_paragraph()
            run_q = p.add_run(f"\u201c{quote}\u201d")
            run_q.italic = True
            run_q.font.size = Pt(10.5)

    # --- Footer note ---
    doc.add_paragraph("")
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("Generated by Legal Transcriber")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # Save
    output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    return output_path


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(description="Generate transcript .docx")
    parser.add_argument("work_dir", help="Path to transcript work directory")
    parser.add_argument("output_path", help="Path for the output .docx file")
    parser.add_argument("--analysis", help="JSON string with analysis data", default=None)
    args = parser.parse_args()

    analysis = None
    if args.analysis:
        try:
            analysis = json.loads(args.analysis)
        except json.JSONDecodeError:
            # Try reading as a file path
            if os.path.isfile(args.analysis):
                with open(args.analysis) as f:
                    analysis = json.load(f)
            else:
                print(f"WARNING: Could not parse analysis JSON", file=sys.stderr)

    try:
        result = build_document(args.work_dir, args.output_path, analysis)
        print(json.dumps({"status": "ok", "output_path": result}))
    except Exception as e:
        print(json.dumps({"status": "error", "error": str(e)}), file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
